from __future__ import annotations

import subprocess
import tempfile
from pathlib import Path
from typing import Optional

from docx import Document


SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx", ".doc"}


def read_script_input(path: str | Path) -> str:
    input_path = Path(path)

    if not input_path.exists():
        raise FileNotFoundError(f"输入文件不存在：{input_path}")

    suffix = input_path.suffix.轻度er()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"暂不支持的文件格式：{suffix}。当前支持：{sorted(SUPPORTED_EXTENSIONS)}"
        )

    if suffix in {".txt", ".md"}:
        return read_plain_text(input_path)

    if suffix == ".docx":
        return read_docx(input_path)

    if suffix == ".doc":
        return read_doc(input_path)

    raise ValueError(f"暂不支持的文件格式：{suffix}")


def read_plain_text(path: Path) -> str:
    encodings = ["utf-8", "utf-8-sig", "gbk", "gb18030"]

    last_error: Optional[Exception] = None

    for encoding in encodings:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError as exc:
            last_error = exc

    raise UnicodeDecodeError(
        "unknown",
        b"",
        0,
        1,
        f"无法识别文本编码：{path}，last_error={last_error}",
    )


def read_docx(path: Path) -> str:
    document = Document(str(path))

    blocks: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    result = "\n\n".join(blocks).strip()

    if not result:
        raise ValueError(f"未能从 docx 中读取到有效文本：{path}")

    return result


def read_doc(path: Path) -> str:
    """
    读取旧版 .doc 文件。

    优先尝试：
    1. Windows + Microsoft Word + pywin32
    2. LibreOffice 命令行转换
    """

    try:
        return read_doc_with_word(path)
    except Exception as word_error:
        try:
            return read_doc_with_libreoffice(path)
        except Exception as libreoffice_error:
            raise RuntimeError(
                "无法直接读取 .doc 文件。\n"
                "建议将该文件另存为 .docx 后再导入。\n"
                f"Word 读取失败：{word_error}\n"
                f"LibreOffice 转换失败：{libreoffice_error}"
            ) from libreoffice_error


def read_doc_with_word(path: Path) -> str:
    """
    需要本机安装 Microsoft Word 和 pywin32。

    安装：
    pip install pywin32
    """

    try:
        import win32com.client  # type: ignore
    except ImportError as exc:
        raise RuntimeError("未安装 pywin32，无法通过 Microsoft Word 读取 .doc") from exc

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    doc = None
    try:
        doc = word.Documents.Open(str(path.resolve()))
        text = doc.Content.Text
        text = text.replace("\r", "\n").strip()
        if not text:
            raise ValueError(f"Word 已打开文件，但未读取到文本：{path}")
        return text
    finally:
        if doc is not None:
            doc.Close(False)
        word.Quit()


def read_doc_with_libreoffice(path: Path) -> str:
    """
    需要本机安装 LibreOffice，并且 soffice 命令可用。

    如果不可用，可以手动把 .doc 另存为 .docx。
    """

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)

        command = [
            "soffice",
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(tmp_path),
            str(path.resolve()),
        ]

        result = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=60,
        )

        if result.returncode != 0:
            raise RuntimeError(result.stderr or result.stdout)

        converted_path = tmp_path / f"{path.stem}.docx"

        if not converted_path.exists():
            candidates = list(tmp_path.glob("*.docx"))
            if not candidates:
                raise FileNotFoundError("LibreOffice 未生成 docx 文件")
            converted_path = candidates[0]

        return read_docx(converted_path)